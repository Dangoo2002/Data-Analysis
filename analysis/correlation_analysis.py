import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_csv(file_name):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(script_dir, '..', 'data')
    return pd.read_csv(os.path.join(data_dir, file_name))

def pearson_correlation(x, y):
    n = len(x)
    
 
    x_mean = np.mean(x)
    y_mean = np.mean(y)
    
   
    dx = x - x_mean
    dy = y - y_mean
    
  
    dxdy = dx * dy
    
  
    dx_squared = dx ** 2
    dy_squared = dy ** 2
    
  
    sum_dxdy = np.sum(dxdy)
    sum_dx_squared = np.sum(dx_squared)
    sum_dy_squared = np.sum(dy_squared)
    
   
    r = sum_dxdy / np.sqrt(sum_dx_squared * sum_dy_squared)
    
    return r, x_mean, y_mean, dx, dy, dxdy, dx_squared, dy_squared, sum_dxdy, sum_dx_squared, sum_dy_squared

def correlation_analysis(days_active, new_users):
   
    days_active['CREATION_DATE'] = pd.to_datetime(days_active['CREATION_DATE'])
    new_users['WEEK'] = pd.to_datetime(new_users['WEEK'])

    
    days_active['CREATION_DATE'] = days_active['CREATION_DATE'].dt.to_period('W').apply(lambda r: r.start_time)
    new_users['WEEK'] = new_users['WEEK'].dt.to_period('W').apply(lambda r: r.start_time)

    merged_df = pd.merge(days_active, new_users, left_on='CREATION_DATE', right_on='WEEK')

    if 'Days Active' in merged_df.columns and 'NEW_USERS' in merged_df.columns:
        x = merged_df['Days Active'].values
        y = merged_df['NEW_USERS'].values
   
        r, x_mean, y_mean, dx, dy, dxdy, dx_squared, dy_squared, sum_dxdy, sum_dx_squared, sum_dy_squared = pearson_correlation(x, y)

        # Create a Word document
        doc = Document()
        doc.add_heading('Pearson Correlation Analysis', 0)

        doc.add_paragraph('Step 1: Calculate means')
        doc.add_paragraph(f'x̄ (Mean of Days Active) = {x_mean:.4f}')
        doc.add_paragraph(f'ȳ (Mean of New Users) = {y_mean:.4f}')
        doc.add_paragraph(f'x̄ = (Σx) / n = {np.sum(x):.4f} / {len(x)} = {x_mean:.4f}')
        doc.add_paragraph(f'ȳ = (Σy) / n = {np.sum(y):.4f} / {len(y)} = {y_mean:.4f}')

        doc.add_paragraph('Step 2: Calculate deviation scores')
        doc.add_paragraph('dx = x - x̄')
        doc.add_paragraph('dy = y - ȳ')
        doc.add_paragraph(f'First few values of dx: {dx[:5]}')
        doc.add_paragraph(f'First few values of dy: {dy[:5]}')

        doc.add_paragraph('Step 3: Calculate product of deviation scores')
        doc.add_paragraph('dxdy = dx * dy')
        doc.add_paragraph(f'First few values of dxdy: {dxdy[:5]}')

        doc.add_paragraph('Step 4: Square deviation scores')
        doc.add_paragraph('dx² = dx * dx')
        doc.add_paragraph('dy² = dy * dy')
        doc.add_paragraph(f'First few values of dx²: {dx_squared[:5]}')
        doc.add_paragraph(f'First few values of dy²: {dy_squared[:5]}')

        doc.add_paragraph('Step 5: Sum all calculated values')
        doc.add_paragraph(f'Σdxdy = {sum_dxdy:.4f}')
        doc.add_paragraph(f'Σdx² = {sum_dx_squared:.4f}')
        doc.add_paragraph(f'Σdy² = {sum_dy_squared:.4f}')

        doc.add_paragraph('Step 6: Calculate correlation coefficient')
        doc.add_paragraph('r = Σdxdy / √(Σdx² * Σdy²)')
        doc.add_paragraph(f'r = {sum_dxdy:.4f} / √({sum_dx_squared:.4f} * {sum_dy_squared:.4f})')
        doc.add_paragraph(f'r = {sum_dxdy:.4f} / {np.sqrt(sum_dx_squared * sum_dy_squared):.4f}')
        doc.add_paragraph(f'r = {r:.6f}')

       
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'output')
        os.makedirs(output_dir, exist_ok=True)

    
        plt.figure(figsize=(10, 6))
        plt.scatter(x, y, alpha=0.5)
        plt.xlabel('Days Active')
        plt.ylabel('New Users')
        plt.title(f'Scatter Plot: Days Active vs New Users\nCorrelation: {r:.6f}')
        scatter_path = os.path.join(output_dir, 'correlation_scatter_plot.png')
        plt.savefig(scatter_path)
        plt.close()

    
        doc.add_picture(scatter_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        plt.figure(figsize=(8, 6))
        correlation_matrix = merged_df[['Days Active', 'NEW_USERS']].corr()
        sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', vmin=-1, vmax=1, center=0)
        plt.title('Correlation Heatmap')
        heatmap_path = os.path.join(output_dir, 'correlation_heatmap.png')
        plt.savefig(heatmap_path)
        plt.close()

        
        doc.add_picture(heatmap_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save the Word document
        doc.save(os.path.join(output_dir, 'correlation_analysis.docx'))

        print("Correlation analysis complete. Results saved in 'output/correlation_analysis.docx'.")
    else:
        print("The required columns for correlation were not found in the datasets.")

if __name__ == "__main__":

    days_active = load_csv('weekly_days_active.csv')
    new_users = load_csv('weekly_new_users.csv')


    correlation_analysis(days_active, new_users)